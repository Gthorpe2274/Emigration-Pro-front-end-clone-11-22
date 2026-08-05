from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "reference.docx"
OUTPUT = ROOT / "EmigrationPro_Affiliate_Outreach_Clean.docx"
LOGO = ROOT.parents[1] / "public" / "images" / "logo-full.png"

BLUE = RGBColor(0x1A, 0x5C, 0x8A)
GRAY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x22, 0x22, 0x22)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    ppr = p.pPr
    if ppr is not None and ppr.numPr is not None:
        ppr.remove(ppr.numPr)
    paragraph.style = "Normal"


def reset_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.first_line_indent = None
    fmt.keep_with_next = False
    fmt.keep_together = False
    fmt.page_break_before = False


def add_text(paragraph, text, size=10.5, bold=False, italic=False, color=BLACK):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def copy_numbering(source_paragraph, target_paragraph):
    source_ppr = source_paragraph._p.pPr
    if source_ppr is None or source_ppr.numPr is None:
        return
    target_ppr = target_paragraph._p.get_or_add_pPr()
    if target_ppr.numPr is not None:
        target_ppr.remove(target_ppr.numPr)
    target_ppr.append(deepcopy(source_ppr.numPr))


doc = Document(SOURCE)
paragraphs = doc.paragraphs
if len(paragraphs) < 53 or not any(run._r.xpath(".//w:br[@w:type='page']") for run in paragraphs[18].runs):
    raise RuntimeError("The reference document structure does not match the audited template.")

# Reuse only the existing cover paragraph slots. Paragraph 18's manual page break
# and every paragraph after it remain untouched.
for paragraph in paragraphs[:18]:
    clear_paragraph(paragraph)

# 0 — logo
p = paragraphs[0]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=7)
p.add_run().add_picture(str(LOGO), width=Inches(3.25))

# 1 — eyebrow
p = paragraphs[1]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=7)
add_text(p, "AFFILIATE PARTNER INVITATION", size=10, bold=True, color=BLUE)

# 2 — subtle spacer
p = paragraphs[2]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=2)

# 3 — headline
p = paragraphs[3]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=9, line=1.05)
add_text(p, "Help Your Audience Plan Their Move Abroad — and Earn 35%–50%", size=23, bold=True, color=BLUE)

# 4 — congratulatory lead
p = paragraphs[4]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=12, line=1.12)
add_text(
    p,
    "Congratulations on building an audience that trusts your guidance about moving abroad. "
    "Emigration Pro helps turn that interest into a practical, personalized relocation plan.",
    size=11.5,
    color=GRAY,
)

# 5 — exclusivity and differentiation
p = paragraphs[5]
reset_format(p, after=10, line=1.08)
add_text(
    p,
    "We are inviting your channel to consider a private affiliate partnership. Unlike generic country guides, "
    "Emigration Pro provides personalized guidance at both the country and city level.",
    size=10.5,
)

# 6 — viewer heading
p = paragraphs[6]
reset_format(p, before=3, after=4)
add_text(p, "What Your Viewers Receive", size=12, bold=True, color=BLUE)

# 7-10 — viewer bullets, using the same Word-native numbering definition as the letter
bullet_texts = [
    "A free personalized city-level assessment",
    "A free destination-specific relocation portal",
    "A comprehensive 14-section Emigration Report",
    "Research supported by government and other reliable sources",
]
letter_bullet = paragraphs[30]
for index, text in enumerate(bullet_texts, start=7):
    p = paragraphs[index]
    p.style = "List Paragraph"
    reset_format(p, after=3, line=1.0)
    copy_numbering(letter_bullet, p)
    add_text(p, text, size=10.5)

# 11 — creator heading
p = paragraphs[11]
reset_format(p, before=6, after=4)
add_text(p, "What You Receive", size=12, bold=True, color=BLUE)

# 12 — commission centerpiece
p = paragraphs[12]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=3)
add_text(p, "35%–50% commission", size=18, bold=True, color=BLUE)
add_text(p, " on every report purchased through your unique affiliate link", size=10.5, color=GRAY)

# 13 — zero-cost support line
p = paragraphs[13]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=10)
add_text(p, "No membership fees. No upfront investment. No overhead.", size=10.5, bold=True, color=GRAY)

# 14 — brand promise
p = paragraphs[14]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8)
add_text(p, "You inspire the move. Emigration Pro helps them plan it.", size=12.5, bold=True, italic=True, color=BLUE)

# 15 — invitation note
p = paragraphs[15]
reset_format(p, after=8, line=1.05)
add_text(
    p,
    "This invitation is being extended to a limited number of carefully selected creators whose audiences closely "
    "align with our service. We are beginning with your channel.",
    size=10,
    color=GRAY,
)

# 16 — CTA
p = paragraphs[16]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, before=2, after=5)
add_text(p, "Review the enclosed partnership proposal to learn more.", size=10.5, bold=True)

# 17 — contact and signature
p = paragraphs[17]
reset_format(p, WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.05)
add_text(p, "EmigrationPro.com", size=11, bold=True, color=BLUE)
add_text(p, "   |   info@emigrationpro.com", size=10, color=GRAY)
add_text(p, "\nGreg Thorpe, Manager  |  The Emigration Pro Team", size=9.5, color=GRAY)

# Improve accessibility metadata for the inserted logo.
drawing = paragraphs[0]._p.xpath(".//wp:docPr")
if drawing:
    drawing[0].set("name", "Emigration Pro logo")
    drawing[0].set("descr", "Emigration Pro — Your Global Journey Starts Here")

doc.core_properties.title = "Emigration Pro Affiliate Partner Outreach"
doc.core_properties.subject = "Affiliate partnership invitation for YouTube creators"
doc.core_properties.author = "The Emigration Pro Team"
doc.core_properties.keywords = "Emigration Pro, affiliate partnership, creator outreach"

doc.save(OUTPUT)
print(OUTPUT)
