from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(__file__).with_name("reference.docx")
doc = Document(path)

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for i, p in enumerate(doc.paragraphs):
    ppr = p._p.pPr
    page_break_before = bool(ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None)
    breaks = sum(1 for br in p._p.iter(qn("w:br")) if br.get(qn("w:type")) == "page")
    numpr = ppr.find(qn("w:numPr")) if ppr is not None else None
    runs = []
    for r in p.runs:
        color = r.font.color.rgb
        runs.append({
            "text": r.text,
            "font": r.font.name,
            "size": r.font.size.pt if r.font.size else None,
            "bold": r.bold,
            "italic": r.italic,
            "color": str(color) if color else None,
        })
    print({
        "i": i,
        "style": p.style.name,
        "text": p.text,
        "alignment": str(p.alignment),
        "before": p.paragraph_format.space_before.pt if p.paragraph_format.space_before else None,
        "after": p.paragraph_format.space_after.pt if p.paragraph_format.space_after else None,
        "line": p.paragraph_format.line_spacing,
        "left_indent": p.paragraph_format.left_indent.inches if p.paragraph_format.left_indent else None,
        "first_indent": p.paragraph_format.first_line_indent.inches if p.paragraph_format.first_line_indent else None,
        "page_break_before": page_break_before,
        "page_breaks": breaks,
        "numbered": numpr is not None,
        "runs": runs,
    })

for i, section in enumerate(doc.sections):
    print("SECTION", i, {
        "page_width": section.page_width.inches,
        "page_height": section.page_height.inches,
        "top": section.top_margin.inches,
        "bottom": section.bottom_margin.inches,
        "left": section.left_margin.inches,
        "right": section.right_margin.inches,
        "header_distance": section.header_distance.inches,
        "footer_distance": section.footer_distance.inches,
    })
